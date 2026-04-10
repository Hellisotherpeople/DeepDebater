"""Export debate transcripts to DOCX format."""

from __future__ import annotations

import os
import re
from html.parser import HTMLParser

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class _HTMLToDocx(HTMLParser):
    """Simple HTML-to-docx converter for debate transcripts."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._current_paragraph = None
        self._bold = False
        self._underline = False
        self._italic = False
        self._in_heading = 0  # 0 = no heading, 1-4 = heading level
        self._in_list = False
        self._list_items: list[str] = []
        self._text_buffer = ""

    def _flush_text(self):
        if self._text_buffer and self._current_paragraph is not None:
            run = self._current_paragraph.add_run(self._text_buffer)
            run.bold = self._bold
            run.underline = self._underline
            run.italic = self._italic
            if self._bold and self._underline:
                run.font.color.rgb = RGBColor(0, 0, 139)  # dark blue for emphasis
            self._text_buffer = ""

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4"):
            self._flush_text()
            level = int(tag[1])
            self._in_heading = level
            self._current_paragraph = self.doc.add_heading("", level=min(level, 4))
        elif tag == "p":
            self._flush_text()
            self._current_paragraph = self.doc.add_paragraph()
        elif tag == "div":
            pass  # don't create new paragraph for divs
        elif tag in ("b", "strong"):
            self._flush_text()
            self._bold = True
        elif tag == "u":
            self._flush_text()
            self._underline = True
        elif tag in ("i", "em"):
            self._flush_text()
            self._italic = True
        elif tag == "mark":
            self._flush_text()
            self._bold = True
            self._underline = True
        elif tag == "br":
            if self._current_paragraph is not None:
                self._flush_text()
                self._current_paragraph.add_run("\n")
        elif tag in ("ul", "ol"):
            self._flush_text()
            self._in_list = True
        elif tag == "li":
            self._flush_text()
            self._current_paragraph = self.doc.add_paragraph(style="List Bullet")
        elif tag == "hr":
            self._flush_text()
            self._current_paragraph = self.doc.add_paragraph()
            self._current_paragraph.add_run("─" * 50)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("h1", "h2", "h3", "h4"):
            self._flush_text()
            self._in_heading = 0
        elif tag in ("b", "strong", "mark"):
            self._flush_text()
            self._bold = False
            if tag == "mark":
                self._underline = False
        elif tag == "u":
            self._flush_text()
            self._underline = False
        elif tag in ("i", "em"):
            self._flush_text()
            self._italic = False
        elif tag in ("ul", "ol"):
            self._in_list = False
        elif tag == "p":
            self._flush_text()
        elif tag == "div":
            self._flush_text()

    def handle_data(self, data):
        text = data
        if not text.strip():
            if text and self._current_paragraph is not None:
                self._text_buffer += " "
            return
        if self._current_paragraph is None:
            self._current_paragraph = self.doc.add_paragraph()
        self._text_buffer += text


def export_debate_to_docx(
    transcript_html: str,
    output_path: str,
    title: str = "Policy Debate Transcript",
) -> str:
    """Convert a debate HTML transcript to a formatted DOCX file.

    Args:
        transcript_html: The full debate transcript as HTML.
        output_path: Where to save the .docx file.
        title: Document title.

    Returns:
        The path to the saved file.
    """
    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # Parse HTML into docx
    parser = _HTMLToDocx(doc)
    parser.feed(transcript_html)
    parser._flush_text()

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def export_speech_to_docx(
    speech_name: str,
    speech_html: str,
    output_path: str,
) -> str:
    """Export a single speech to a DOCX file."""
    return export_debate_to_docx(speech_html, output_path, title=speech_name)
